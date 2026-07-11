import io

import pytest
from docx import Document
from fpdf import FPDF

from app.ingest import extract_text


def test_pasted_text_passthrough():
    assert extract_text(None, None, "  hello world  ", 1000) == "hello world"


def test_truncates_to_max_chars():
    assert extract_text(None, None, "a" * 50, 10) == "a" * 10


def test_docx_extraction():
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("See Brown v. Board, 347 U.S. 483 (1954).")
    buf = io.BytesIO()
    doc.save(buf)
    text = extract_text("brief.docx", buf.getvalue(), None, 10000)
    assert "347 U.S. 483" in text
    assert "First paragraph." in text


def test_pdf_extraction():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(text="See Miranda v. Arizona, 384 U.S. 436 (1966).")
    data = bytes(pdf.output())
    text = extract_text("brief.pdf", data, None, 10000)
    assert "384 U.S. 436" in text


def test_empty_raises():
    with pytest.raises(ValueError):
        extract_text(None, None, "   ", 1000)
